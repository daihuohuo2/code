from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"C:\Users\20694\Desktop\code\信息安全作业一.docx"


def set_east_asian_font(run, font_name="宋体"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(10.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_east_asian_font(run, "黑体")
    run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_para(doc, text="", first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_east_asian_font(run)
        run.font.size = Pt(10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_east_asian_font(run)
        run.font.size = Pt(10.5)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_run = header.add_run("信息安全作业一")
    set_east_asian_font(h_run, "宋体")
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = RGBColor(100, 100, 100)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run("第 ")
    set_east_asian_font(f_run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    f_run._r.append(fld_begin)
    f_run._r.append(instr)
    f_run._r.append(fld_end)
    f_run2 = footer.add_run(" 页")
    set_east_asian_font(f_run2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("信息安全作业一")
    set_east_asian_font(r, "黑体")
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)

    meta = doc.add_table(rows=3, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    data = [
        ("课程", "信息安全", "作业", "作业一"),
        ("姓名", "__________", "学号", "__________"),
        ("日期", "2026 年 5 月", "主题", "恶意代码、漏洞与安全态势分析"),
    ]
    for row, values in zip(meta.rows, data):
        for idx, value in enumerate(values):
            set_cell_text(row.cells[idx], value, bold=(idx in (0, 2)))
            if idx in (0, 2):
                set_cell_shading(row.cells[idx], "EAF2F8")

    add_para(doc, "")

    add_heading(doc, "一、几种典型的恶意代码传播方式及防范措施", 1)
    add_para(doc, "恶意代码是指未经授权、以破坏系统、窃取信息、控制主机或牟利为目的的程序或脚本。它并不只通过单一方式传播，而是常常结合社会工程学、系统漏洞和弱口令等手段形成攻击链。常见传播方式如下。")
    add_bullets(doc, [
        "电子邮件和即时通信传播：攻击者将木马、宏病毒或钓鱼链接伪装成通知、发票、作业资料、压缩包等，诱导用户点击或下载。",
        "网页挂马和恶意链接传播：用户访问被入侵的网站或钓鱼站点时，浏览器、插件或脚本漏洞可能被利用，恶意代码被自动下载执行。",
        "移动存储介质传播：U 盘、移动硬盘等设备可能携带自动运行脚本或伪装文件，在多台电脑之间交叉感染。",
        "漏洞利用传播：蠕虫利用操作系统、数据库、中间件、Web 应用或网络服务漏洞自动扫描并感染其他主机。",
        "软件供应链传播：攻击者污染安装包、更新通道、第三方依赖库或开源组件，使用户在正常安装和升级时引入恶意代码。",
        "弱口令和远程服务传播：通过爆破 RDP、SSH、数据库、NAS、路由器后台等入口获得权限，再植入木马或挖矿程序。",
    ])
    add_para(doc, "防范恶意代码传播应坚持“人、终端、网络、应用、数据”协同治理。个人层面要提高安全意识，不随意打开陌生附件，不安装来源不明的软件，不把 U 盘在多台设备之间无检查地混用。终端层面应及时更新操作系统和应用补丁，开启杀毒与主机防护，关闭不必要的自动运行功能。网络层面要限制高危端口暴露，采用防火墙、入侵检测、邮件网关和 Web 安全网关。管理层面要落实最小权限、强口令、多因素认证、备份恢复和日志审计。对于学校和单位，还应定期开展漏洞扫描、钓鱼演练和应急演练，避免一次点击或一个弱口令演变成大范围感染。")

    add_heading(doc, "二、什么是僵尸程序和僵尸网络，有何危害", 1)
    add_para(doc, "僵尸程序通常指被攻击者植入受害主机中的远程控制程序。受害设备表面上仍能正常使用，但实际已经被攻击者控制，可以接收命令、执行下载、扫描、攻击、窃密、代理转发等操作。被控制的主机常被称为“僵尸主机”。")
    add_para(doc, "僵尸网络是大量僵尸主机按照攻击者指令组成的控制网络。攻击者通过命令控制服务器、P2P 控制结构、域名生成算法或社交平台等方式下发指令，使分散在不同地区的设备协同实施攻击。其危害主要体现在以下方面：")
    add_bullets(doc, [
        "发动 DDoS 攻击，利用大量主机同时访问目标网站或服务器，造成服务瘫痪。",
        "发送垃圾邮件、钓鱼邮件和恶意链接，扩大诈骗、木马和勒索软件传播范围。",
        "窃取账号、浏览器 Cookie、文件、键盘记录和隐私数据，导致个人和单位信息泄露。",
        "作为攻击跳板隐藏攻击者真实来源，使溯源和取证更困难。",
        "控制物联网摄像头、路由器等设备，形成长期潜伏的网络基础设施威胁。",
        "消耗受害设备资源，导致电脑卡顿、流量异常、电费增加，甚至被用于挖矿牟利。",
    ])
    add_para(doc, "因此，治理僵尸网络既需要用户加固终端，也需要运营商、安全机构、厂商和监管部门协同处置，包括封堵控制域名、清理恶意样本、通报受感染主机、修复漏洞和提升默认口令安全。")

    add_heading(doc, "三、典型漏洞、产生原因及防护思路", 1)
    add_para(doc, "漏洞是信息系统在设计、编码、配置、部署或运维过程中形成的安全缺陷。典型漏洞及原因如下。")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["漏洞类型", "产生原因", "可能危害"]):
        set_cell_text(cell, text, bold=True)
        set_cell_shading(cell, "D9EAF7")
    rows = [
        ("SQL 注入", "程序把用户输入直接拼接进 SQL 语句，缺少参数化查询和输入校验。", "绕过登录、拖库、篡改成绩或业务数据。"),
        ("跨站脚本 XSS", "页面未对用户输入和输出进行转义，攻击脚本被浏览器执行。", "盗取 Cookie、劫持会话、冒充用户操作。"),
        ("文件上传漏洞", "只检查文件后缀或 Content-Type，未校验内容、权限和存储路径。", "上传 WebShell，获得服务器控制权。"),
        ("越权访问", "只在前端隐藏按钮，后端没有检查用户身份和资源归属。", "普通用户查看或修改他人数据。"),
        ("缓冲区溢出", "C/C++ 程序未检查边界，写入超过缓冲区长度。", "程序崩溃、远程代码执行、提权。"),
        ("弱口令与默认口令", "账号策略薄弱，设备或系统上线后未修改默认密码。", "被爆破登录，成为入侵入口。"),
        ("配置错误", "调试接口、目录列表、敏感端口或云存储权限暴露。", "源代码、密钥、数据库或个人信息泄露。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            set_cell_text(cell, text)

    add_para(doc, "漏洞很难被完全杜绝。原因在于软件规模越来越大，依赖组件越来越复杂，业务需求不断变化，开发人员认知和测试覆盖也存在边界。即使某个系统当前没有已知漏洞，也可能在新功能上线、依赖升级、配置变更或攻击技术进步后暴露新风险。因此，安全目标不应是“永远没有漏洞”，而应是“尽早发现、快速修复、降低利用成功率和危害范围”。")
    add_para(doc, "防止漏洞造成危害，需要建立全生命周期防护：开发阶段进行安全需求分析、代码审计、参数化查询、权限模型设计和依赖组件管理；测试阶段开展漏洞扫描、渗透测试和模糊测试；部署阶段最小化暴露面，关闭默认账号和不必要端口，启用 HTTPS、WAF、主机防护和日志审计；运维阶段建立漏洞通报、补丁评估、灰度升级、备份恢复和应急响应机制。对高危漏洞，必须做到定级、确认影响范围、临时缓解、补丁修复和复盘闭环。")

    add_heading(doc, "四、校园网信息系统可能存在的漏洞和威胁", 1)
    add_para(doc, "以教务管理系统、学生成绩查询系统为例，其安全目标包括身份真实性、数据保密性、成绩完整性、服务可用性和操作可追溯性。校园网系统用户多、角色复杂、访问时间集中，既面临外部攻击，也面临内部越权和误操作风险。")
    add_bullets(doc, [
        "身份认证风险：弱口令、学号作为默认密码、验证码缺陷、找回密码问题过于简单、会话超时设置不合理。",
        "权限控制风险：学生越权查看他人成绩，教师越权修改非本人课程成绩，管理员接口缺少二次确认。",
        "Web 应用漏洞：SQL 注入、XSS、CSRF、文件上传、路径遍历、接口未鉴权等。",
        "数据泄露风险：导出的 Excel 成绩表、身份证号、手机号、家庭住址等敏感数据未脱敏或被误发。",
        "接口与移动端风险：App、小程序、统一身份认证接口和第三方系统对接时，Token 泄露或接口签名校验不足。",
        "运维配置风险：测试环境连接真实数据库，后台管理入口暴露公网，服务器补丁滞后，日志中记录明文密码。",
        "可用性威胁：选课、查分、缴费等高峰期可能遭遇拥塞、爬虫、恶意刷接口或 DDoS 攻击。",
        "内部人员风险：管理员账号共享、离职人员权限未回收、教师电脑中毒后成绩管理账号被盗用。",
    ])
    add_para(doc, "针对校园网系统，建议采用统一身份认证和多因素认证，强制修改默认密码；按角色、课程、学院和数据归属实施细粒度授权；对成绩修改、批量导出、管理员登录等敏感操作进行日志审计和二次确认；对个人信息进行脱敏展示；定期开展代码审计、漏洞扫描和渗透测试；关键数据库实施备份、容灾和最小权限访问。学生和教师端也要加强安全教育，避免账号外借和在不可信设备上登录。")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "五、阅读报告：我国互联网网络安全态势与漏洞治理思考", 1)
    add_para(doc, "通过阅读国家互联网应急中心、国家信息安全漏洞共享平台等公开材料，可以看到当前网络安全威胁已经从单点攻击转向体系化、自动化和持续化攻击。攻击者不再只依赖某一种木马或某一个漏洞，而是会把钓鱼邮件、弱口令爆破、漏洞利用、权限提升、横向移动、数据窃取和勒索破坏串联起来。对于学校、企业和个人而言，网络安全不再只是安装杀毒软件的问题，而是涉及资产管理、漏洞治理、身份认证、数据保护、监测预警和应急响应的综合能力。")
    add_para(doc, "首先，恶意代码传播仍然是重要威胁。随着办公系统、网盘、即时通信和移动终端被广泛使用，攻击者更容易利用“熟人关系”和“业务场景”降低用户警惕。例如伪装成通知文件、课程资料、报销材料、会议链接的恶意附件，往往比单纯的技术攻击更容易成功。一旦终端执行恶意代码，攻击者就可能窃取浏览器凭据、远程控制电脑，或者继续向通讯录联系人传播。因此，安全防护不能只依赖边界防火墙，还要关注终端行为检测、邮件过滤、宏脚本控制、下载文件信誉判断和用户安全意识培训。")
    add_para(doc, "其次，漏洞风险呈现数量多、利用快、影响广的特点。CNVD 等漏洞平台的意义在于汇聚厂商、安全机构、运营单位和研究人员的漏洞信息，使漏洞能够被统一编号、验证、通报和处置。对防守方来说，漏洞信息不是简单的新闻，而是补丁优先级排序的依据。一个系统是否安全，不仅取决于是否存在漏洞，更取决于漏洞暴露面有多大、是否已经出现利用代码、系统中是否保存敏感数据、是否存在可替代的缓解措施。比如同样是 Web 漏洞，暴露在公网的教务系统后台显然比内网测试页面风险更高；同样是组件漏洞，直接处理用户上传文件的组件比离线工具更应优先修复。")
    add_para(doc, "再次，僵尸网络和 DDoS 攻击说明大量普通设备也可能成为公共网络安全风险的一部分。摄像头、路由器、NAS、打印机等设备如果使用默认口令、固件长期不更新，就可能被攻击者批量控制。对个人来说，设备被控制会造成隐私泄露和网络异常；对社会来说，大量设备组成僵尸网络后可以攻击网站、游戏平台、学校服务平台甚至重要行业系统。因此，物联网设备也应纳入资产管理，定期修改默认口令、关闭不必要的远程访问、更新固件，并通过网络侧监测异常流量。")
    add_para(doc, "结合校园场景，信息系统最值得关注的是身份、权限和数据。教务系统、成绩系统和统一认证平台承载大量学生个人信息和教学管理数据，一旦被攻击，可能导致成绩被篡改、个人信息泄露、账号被冒用，甚至影响学校正常教学秩序。校园系统用户基数大，学生安全意识差异明显，密码复用普遍存在，攻击者可以通过撞库、钓鱼和弱口令爆破获取入口。系统建设时应把“默认不信任”作为原则：前端隐藏按钮不能替代后端鉴权；登录成功不能代表可以访问所有资源；教师、学生、管理员必须按最小权限区分；成绩修改、批量导出、权限变更等操作必须可审计、可追溯、可回滚。")
    add_para(doc, "从治理角度看，漏洞无法完全避免，但可以通过制度和技术降低风险。第一，要建立资产清单，知道自己有哪些服务器、域名、系统、接口、数据库和第三方组件。没有资产清单，就无法判断漏洞影响范围。第二，要建立补丁和变更流程，对高危漏洞快速处置，对一般漏洞按周期修复，同时避免补丁造成业务中断。第三，要建立监测能力，关注异常登录、异常导出、短时间大量请求、失败登录暴增、WebShell 文件落地等信号。第四，要建立应急预案，明确发现入侵后如何隔离主机、保护现场、备份日志、重置凭据、通知用户和恢复业务。第五，要开展安全教育，让普通用户知道钓鱼邮件、恶意附件、弱口令和账号外借的危害。")
    add_para(doc, "总体来看，网络安全态势的核心变化是攻击门槛降低而防守复杂度升高。自动化扫描工具、公开漏洞利用代码和黑灰产分工，使攻击者可以更快发现和利用薄弱系统；而防守方必须同时面对人员、技术、流程和管理问题。阅读相关报告后，我认为安全建设不能追求一次性“做完”，而要形成持续改进机制。对个人来说，要养成更新系统、使用强密码、开启多因素认证、谨慎打开附件和备份重要数据的习惯。对学校和单位来说，要把漏洞管理、日志审计、权限治理和应急响应纳入日常工作。只有把安全从“出事后补救”转变为“事前预防、事中发现、事后复盘”，才能真正降低恶意代码、漏洞利用和数据泄露带来的危害。")

    add_heading(doc, "参考资料", 1)
    add_numbered(doc, [
        "国家互联网应急中心（CNCERT）：https://www.cert.org.cn/publish/main/17/index.html",
        "国家信息安全漏洞共享平台（CNVD）：https://www.cnvd.org.cn/",
        "CNCERT 互联网安全威胁报告与相关公开材料：https://www.cert.org.cn/",
        "CNVD 漏洞周报与漏洞通报公开信息：https://www.cnvd.org.cn/flaw/list",
    ])

    doc.save(OUT)


if __name__ == "__main__":
    build()
